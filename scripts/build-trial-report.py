import json, re, sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

data=json.loads(Path(sys.argv[1]).read_text(encoding="utf-8")); out=Path(sys.argv[2])
project=re.search(r"项目名称：\s*([^\n]+)",data["first"]["text"]).group(1).strip()
owner="如皋市公安局"; contractor="中通服咨询设计研究院有限公司"; supervisor="苏州市软件评测中心有限公司"
def grab(text,label):
    m=re.search(label+r"：([^\n]+)",text); return m.group(1).strip() if m else ""
issues=[{"date":x["date"],"problem":grab(x["text"],"存在问题"),"problem_time":grab(x["text"],"问题发生时间"),"repair_time":grab(x["text"],"修复时间"),"repair":grab(x["text"],"修复情况")} for x in data["abnormal"]]

doc=Document(); sec=doc.sections[0]; sec.page_width=Inches(8.27); sec.page_height=Inches(11.69); sec.top_margin=Inches(.8); sec.bottom_margin=Inches(.75); sec.left_margin=Inches(.82); sec.right_margin=Inches(.82); sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)
styles=doc.styles
normal=styles["Normal"]; normal.font.name="宋体"; normal._element.rPr.rFonts.set(qn("w:eastAsia"),"宋体"); normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.35
for name,size,color,before,after in [("Heading 1",16,"1F4E79",14,7),("Heading 2",13,"1F4E79",10,5)]:
    s=styles[name]; s.font.name="黑体"; s._element.rPr.rFonts.set(qn("w:eastAsia"),"黑体"); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
header=sec.header.paragraphs[0]; header.text="如皋市新时代技防城项目｜系统试运行报告"; header.style=normal; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; header.runs[0].font.size=Pt(9); header.runs[0].font.color.rgb=RGBColor(100,110,120)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=footer.add_run("第 "); fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); run._r.addnext(fld); footer.add_run(" 页")

for _ in range(3): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("系统试运行报告"); r.bold=True; r.font.name="黑体"; r._element.rPr.rFonts.set(qn("w:eastAsia"),"黑体"); r.font.size=Pt(30); r.font.color.rgb=RGBColor(31,78,121)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(project); r.font.name="黑体"; r._element.rPr.rFonts.set(qn("w:eastAsia"),"黑体"); r.font.size=Pt(16)
doc.add_paragraph()
meta=doc.add_table(rows=5,cols=2); meta.alignment=WD_TABLE_ALIGNMENT.CENTER; meta.autofit=False
rows=[("建设单位",owner),("承建单位",contractor),("监理单位",supervisor),("试运行周期","2025年11月13日至2026年8月7日"),("报告日期","2026年8月12日")]
for i,(a,b) in enumerate(rows): meta.cell(i,0).width=Inches(1.4);meta.cell(i,1).width=Inches(4.8);meta.cell(i,0).text=a;meta.cell(i,1).text=b
for row in meta.rows:
    for j,c in enumerate(row.cells): c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER;c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if j==0 else WD_ALIGN_PARAGRAPH.LEFT
doc.add_page_break()

doc.add_heading("一、项目概况",level=1); doc.add_paragraph(f"{project}完成建设后，于2025年11月13日至2026年8月7日开展系统试运行。本次报告依据压缩包内189份逐日《系统试运行记录》汇总形成，用于说明试运行组织、检查内容、运行表现、问题处置及综合结论。")
doc.add_heading("二、试运行组织与范围",level=1)
for t in ["建设单位：如皋市公安局。","承建单位：中通服咨询设计研究院有限公司。","监理单位：苏州市软件评测中心有限公司。","试运行记录覆盖189个工作日，检查范围包括图像网、公安网、服务器运行及运维治理。"]: doc.add_paragraph(t,style="List Bullet")
doc.add_heading("三、试运行检查内容",level=1)
checks=[("图像网","抽查不同组织目录点位的视频预览与回放取流；核查人脸、人体、车辆数据查询；验证人脸搜图及大数据后端服务。"),("公安网","抽查视频预览与回放；核查智能搜索、历史布控预警、人员档案、全息档案、关系图谱及历史建模任务。"),("服务器","检查图像网通用服务器、云存储服务器、解析服务器、大数据服务器，以及公安网通用服务器和大数据服务器运行状态。"),("运维治理","持续检查治理任务执行情况，逐日记录治理任务完成情况。")]
for title,desc in checks: doc.add_heading(title,level=2); doc.add_paragraph(desc)
doc.add_heading("四、试运行结果统计",level=1)
table=doc.add_table(rows=1,cols=4); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.style="Table Grid"; hdr=table.rows[0].cells
for c,t in zip(hdr,["记录总数","正常记录","异常记录","异常闭环率"]): c.text=t
row=table.add_row().cells
for c,t in zip(row,["189份","178份","11份","100%（11/11）"]): c.text=t
for row in table.rows:
    for c in row.cells: c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER;c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("总体情况：试运行期间各项日常检查持续开展。全部11项异常均记录了问题发生时间、修复时间及处理结果，未发现记录中存在未闭环事项。")
doc.add_heading("五、异常问题及闭环情况",level=1)
t=doc.add_table(rows=1,cols=5);t.style="Table Grid";t.alignment=WD_TABLE_ALIGNMENT.CENTER
for c,v in zip(t.rows[0].cells,["日期","发生/修复时间","存在问题","修复措施及结果","状态"]):c.text=v
for x in issues:
    cells=t.add_row().cells; times=f'{x["problem_time"].split()[-1]} / {x["repair_time"].split()[-1]}'
    for c,v in zip(cells,[x["date"],times,x["problem"],x["repair"],"已闭环"]):c.text=v
for i,row in enumerate(t.rows):
    for j,c in enumerate(row.cells): c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER;c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER if j in (0,1,4) else WD_ALIGN_PARAGRAPH.LEFT
doc.add_paragraph("注：以上异常明细、发生时间、修复时间和处理结果均摘自对应日期的试运行记录。")
doc.add_heading("六、综合评价",level=1)
for text in ["系统功能方面：图像网、公安网相关视频、检索、布控、档案、图谱及建模等功能在逐日检查记录中总体运行正常。","基础设施方面：图像网和公安网相关服务器总体运行稳定，出现的授权、服务、配置、接口及数据类问题均完成处置。","运维保障方面：试运行期间形成了逐日检查和异常闭环记录，问题响应及恢复过程可追溯。"]: doc.add_paragraph(text,style="List Bullet")
doc.add_heading("七、试运行结论",level=1)
p=doc.add_paragraph(); r=p.add_run("经汇总分析，系统在2025年11月13日至2026年8月7日试运行期间总体运行稳定，主要功能和服务满足试运行要求；发现的11项问题均已完成处理并形成闭环记录。依据现有试运行记录，系统具备进入项目终验及后续正式运行阶段的条件。");r.bold=True
doc.add_paragraph("本报告结论基于压缩包内189份系统试运行记录形成。")
out.parent.mkdir(parents=True,exist_ok=True);doc.save(out);print(out)
